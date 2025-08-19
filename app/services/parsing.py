"""
Enhanced resume parsing with multiple format support.
"""
import logging
import re
from pathlib import Path
from typing import Dict, Any, List, Optional
import tempfile

# Document processing imports
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    import cv2
    import numpy as np
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)

class ResumeParser:
    """Enhanced resume parser with multiple format support."""
    
    def __init__(self):
        """Initialize the resume parser."""
        self.supported_formats = ['.txt', '.docx', '.pdf']
        if OCR_AVAILABLE:
            self.supported_formats.extend(['.png', '.jpg', '.jpeg', '.tiff', '.bmp'])
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a resume file and extract structured information.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing parsed resume data
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        # Extract text based on file type
        text = self._extract_text(file_path)
        
        if not text.strip():
            logger.warning(f"No text extracted from: {file_path}")
            return self._empty_resume_data()
        
        # Parse structured information
        parsed_data = self._parse_text(text)
        parsed_data['file_path'] = str(file_path)
        parsed_data['file_name'] = file_path.name
        parsed_data['raw_text'] = text
        
        return parsed_data
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from various file formats."""
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == '.txt':
                return self._extract_text_txt(file_path)
            elif suffix == '.docx' and DOCX_AVAILABLE:
                return self._extract_text_docx(file_path)
            elif suffix == '.pdf' and PDF_AVAILABLE:
                return self._extract_text_pdf(file_path)
            elif suffix in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'] and OCR_AVAILABLE:
                return self._extract_text_ocr(file_path)
            else:
                logger.error(f"Unsupported file format: {suffix}")
                return ""
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return ""
    
    def _extract_text_txt(self, file_path: Path) -> str:
        """Extract text from plain text files."""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Failed to read text file: {e}")
            return ""
    
    def _extract_text_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return ""
        
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            return ""
    
    def _extract_text_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files."""
        if not PDF_AVAILABLE:
            logger.error("PyPDF2 not available")
            return ""
        
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page: {e}")
                        continue
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            return ""
    
    def _extract_text_ocr(self, file_path: Path) -> str:
        """Extract text from images using OCR."""
        if not OCR_AVAILABLE:
            logger.error("OCR dependencies not available")
            return ""
        
        try:
            # Read image
            img = cv2.imread(str(file_path))
            if img is None:
                logger.error(f"Failed to load image: {file_path}")
                return ""
            
            # Preprocess image for better OCR
            img = self._preprocess_image_for_ocr(img)
            
            # Convert to PIL Image
            img_pil = Image.fromarray(img)
            
            # Perform OCR
            text = pytesseract.image_to_string(img_pil, config='--psm 6')
            return text
        except Exception as e:
            logger.error(f"Failed to extract text using OCR: {e}")
            return ""
    
    def _preprocess_image_for_ocr(self, img: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR results."""
        # Convert to grayscale
        if len(img.shape) == 3:
            img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Apply Gaussian blur to reduce noise
        img = cv2.GaussianBlur(img, (1, 1), 0)
        
        # Apply threshold to get binary image
        _, img = cv2.threshold(img, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return img
    
    def _parse_text(self, text: str) -> Dict[str, Any]:
        """Parse structured information from text."""
        parsed_data = self._empty_resume_data()
        
        # Clean and normalize text
        text = self._clean_text(text)
        
        # Extract different sections
        parsed_data['contact_info'] = self._extract_contact_info(text)
        parsed_data['skills'] = self._extract_skills(text)
        parsed_data['experience'] = self._extract_experience(text)
        parsed_data['education'] = self._extract_education(text)
        parsed_data['certifications'] = self._extract_certifications(text)
        parsed_data['languages'] = self._extract_languages(text)
        
        # Extract summary/objective
        parsed_data['summary'] = self._extract_summary(text)
        
        # Calculate experience years
        parsed_data['total_experience_years'] = self._calculate_experience_years(parsed_data['experience'])
        
        return parsed_data
    
    def _empty_resume_data(self) -> Dict[str, Any]:
        """Return empty resume data structure."""
        return {
            'contact_info': {
                'name': '',
                'email': '',
                'phone': '',
                'location': '',
                'linkedin': '',
                'website': ''
            },
            'skills': [],
            'experience': [],
            'education': [],
            'certifications': [],
            'languages': [],
            'summary': '',
            'total_experience_years': 0,
            'raw_text': '',
            'file_path': '',
            'file_name': ''
        }
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with parsing
        text = re.sub(r'[^\w\s@.\-+(),:;/]', ' ', text)
        
        return text.strip()
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information."""
        contact_info = {
            'name': '',
            'email': '',
            'phone': '',
            'location': '',
            'linkedin': '',
            'website': ''
        }
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Phone pattern (various formats)
        phone_patterns = [
            r'\b(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b',
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
            r'\b\(\d{3}\)\s?\d{3}[-.\s]?\d{4}\b'
        ]
        
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact_info['phone'] = phone_match.group()
                break
        
        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/[A-Za-z0-9\-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group()
        
        # Website pattern
        website_pattern = r'(?:https?://)?(?:www\.)?[A-Za-z0-9\-]+\.[A-Za-z]{2,}(?:/[A-Za-z0-9\-._~:/?#[\]@!$&\'()*+,;=]*)?'
        website_match = re.search(website_pattern, text)
        if website_match and 'linkedin' not in website_match.group().lower():
            contact_info['website'] = website_match.group()
        
        # Name extraction (first few words, typically at the beginning)
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if line and not any(char in line for char in '@+()') and len(line.split()) <= 4:
                # Likely a name if it's short and doesn't contain contact info
                if not re.search(r'\d', line):  # No digits
                    contact_info['name'] = line
                    break
        
        return contact_info
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text."""
        skills = []
        
        # Common skill keywords
        skill_keywords = [
            # Programming languages
            'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift',
            'kotlin', 'scala', 'r', 'matlab', 'sql', 'html', 'css', 'typescript',
            
            # Frameworks and libraries
            'react', 'angular', 'vue', 'django', 'flask', 'spring', 'express', 'nodejs',
            'laravel', 'rails', 'jquery', 'bootstrap', 'tensorflow', 'pytorch', 'pandas',
            'numpy', 'scikit-learn',
            
            # Databases
            'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'sqlite', 'oracle',
            
            # Cloud and DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'linux',
            'ansible', 'terraform',
            
            # Data and Analytics
            'tableau', 'powerbi', 'excel', 'statistics', 'machine learning', 'data analysis',
            'big data', 'hadoop', 'spark',
            
            # Hotel industry specific
            'pms', 'opera', 'fidelio', 'hospitality', 'front desk', 'housekeeping',
            'food service', 'customer service', 'concierge', 'reservations'
        ]
        
        text_lower = text.lower()
        
        # Find skills mentioned in text
        for skill in skill_keywords:
            if skill.lower() in text_lower:
                skills.append(skill.title())
        
        # Look for skills sections
        skills_section_pattern = r'(?:skills?|technical skills?|competencies)[\s:]*([^\n]*(?:\n[^\n]*)*?)(?:\n\n|\n[A-Z]|$)'
        skills_match = re.search(skills_section_pattern, text, re.IGNORECASE | re.MULTILINE)
        
        if skills_match:
            skills_text = skills_match.group(1)
            # Split by common delimiters
            additional_skills = re.split(r'[,;|\n•·-]', skills_text)
            for skill in additional_skills:
                skill = skill.strip()
                if skill and len(skill) < 50:  # Reasonable skill length
                    skills.append(skill)
        
        # Remove duplicates and return
        return list(set(skills))
    
    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience from text."""
        experience = []
        
        # Look for experience section
        exp_patterns = [
            r'(?:work experience|professional experience|employment|career history)[\s:]*([^\n]*(?:\n[^\n]*)*?)(?:\n\n|\n(?:education|skills)|$)',
            r'(?:experience)[\s:]*([^\n]*(?:\n[^\n]*)*?)(?:\n\n|\n(?:education|skills)|$)'
        ]
        
        for pattern in exp_patterns:
            exp_match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if exp_match:
                exp_text = exp_match.group(1)
                
                # Parse individual jobs
                job_entries = self._parse_job_entries(exp_text)
                experience.extend(job_entries)
                break
        
        return experience
    
    def _parse_job_entries(self, exp_text: str) -> List[Dict[str, Any]]:
        """Parse individual job entries from experience text."""
        jobs = []
        
        # Split by likely job separators
        job_blocks = re.split(r'\n(?=\S)', exp_text)
        
        for block in job_blocks:
            block = block.strip()
            if not block:
                continue
            
            job = {
                'title': '',
                'company': '',
                'location': '',
                'start_date': '',
                'end_date': '',
                'description': '',
                'duration_months': 0
            }
            
            lines = block.split('\n')
            
            # First line often contains title and company
            if lines:
                first_line = lines[0].strip()
                # Try to split title and company
                if ' at ' in first_line:
                    parts = first_line.split(' at ', 1)
                    job['title'] = parts[0].strip()
                    job['company'] = parts[1].strip()
                elif ' - ' in first_line:
                    parts = first_line.split(' - ', 1)
                    job['title'] = parts[0].strip()
                    job['company'] = parts[1].strip()
                else:
                    job['title'] = first_line
            
            # Look for dates
            date_pattern = r'(\d{4})\s*[-–]\s*(\d{4}|present|current)'
            date_match = re.search(date_pattern, block, re.IGNORECASE)
            if date_match:
                job['start_date'] = date_match.group(1)
                job['end_date'] = date_match.group(2)
                
                # Calculate duration
                start_year = int(date_match.group(1))
                end_year = 2024 if date_match.group(2).lower() in ['present', 'current'] else int(date_match.group(2))
                job['duration_months'] = (end_year - start_year) * 12
            
            # Rest is description
            if len(lines) > 1:
                job['description'] = '\n'.join(lines[1:]).strip()
            
            if job['title'] or job['company']:  # Only add if we have some info
                jobs.append(job)
        
        return jobs
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education information."""
        education = []
        
        # Look for education section
        edu_pattern = r'(?:education|academic background)[\s:]*([^\n]*(?:\n[^\n]*)*?)(?:\n\n|\n(?:experience|skills)|$)'
        edu_match = re.search(edu_pattern, text, re.IGNORECASE | re.MULTILINE)
        
        if edu_match:
            edu_text = edu_match.group(1)
            
            # Look for degree patterns
            degree_patterns = [
                r'(bachelor|master|phd|doctorate|associate|diploma|certificate).*?(?:in\s+)?(.*?)(?:\n|$)',
                r'(b\.?s\.?|m\.?s\.?|m\.?a\.?|b\.?a\.?|ph\.?d\.?).*?(?:in\s+)?(.*?)(?:\n|$)'
            ]
            
            for pattern in degree_patterns:
                for match in re.finditer(pattern, edu_text, re.IGNORECASE):
                    education.append({
                        'degree': match.group(1).strip(),
                        'field': match.group(2).strip() if match.group(2) else '',
                        'institution': '',
                        'year': ''
                    })
        
        return education
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications."""
        certifications = []
        
        # Look for certification section
        cert_pattern = r'(?:certifications?|licenses?)[\s:]*([^\n]*(?:\n[^\n]*)*?)(?:\n\n|\n[A-Z]|$)'
        cert_match = re.search(cert_pattern, text, re.IGNORECASE | re.MULTILINE)
        
        if cert_match:
            cert_text = cert_match.group(1)
            cert_lines = cert_text.split('\n')
            
            for line in cert_lines:
                line = line.strip()
                if line and len(line) < 100:  # Reasonable cert name length
                    certifications.append(line)
        
        return certifications
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extract languages."""
        languages = []
        
        # Common languages
        language_keywords = [
            'english', 'spanish', 'french', 'german', 'italian', 'portuguese',
            'chinese', 'japanese', 'korean', 'arabic', 'russian', 'hindi'
        ]
        
        text_lower = text.lower()
        
        for lang in language_keywords:
            if lang in text_lower:
                languages.append(lang.title())
        
        return list(set(languages))
    
    def _extract_summary(self, text: str) -> str:
        """Extract summary or objective."""
        # Look for summary section
        summary_patterns = [
            r'(?:summary|objective|profile|about)[\s:]*([^\n]*(?:\n[^\n]*)*?)(?:\n\n|\n[A-Z]|$)',
            r'^([^\n]*(?:\n[^\n]*){0,3})(?:\n\n|$)'  # First few lines
        ]
        
        for pattern in summary_patterns:
            summary_match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if summary_match:
                summary = summary_match.group(1).strip()
                if len(summary) > 50:  # Reasonable summary length
                    return summary
        
        return ''
    
    def _calculate_experience_years(self, experience: List[Dict[str, Any]]) -> float:
        """Calculate total experience years."""
        total_months = 0
        
        for job in experience:
            total_months += job.get('duration_months', 0)
        
        return round(total_months / 12, 1)
    
    def batch_parse(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """Parse multiple resume files."""
        results = []
        
        for file_path in file_paths:
            try:
                parsed_data = self.parse_resume(file_path)
                results.append(parsed_data)
            except Exception as e:
                logger.error(f"Failed to parse {file_path}: {e}")
                # Add empty result with error info
                empty_data = self._empty_resume_data()
                empty_data['file_path'] = file_path
                empty_data['error'] = str(e)
                results.append(empty_data)
        
        return results


# Global parser instance
_resume_parser: Optional[ResumeParser] = None


def get_resume_parser() -> ResumeParser:
    """Get the global resume parser instance."""
    global _resume_parser
    if _resume_parser is None:
        _resume_parser = ResumeParser()
    return _resume_parser
