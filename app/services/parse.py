"""Resume parsing service - simplified without OCR."""
import io
import mimetypes
from typing import Tuple, Optional

import magic
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document


def detect_file_type(file_content: bytes, filename: str) -> str:
    """Detect file type from content and filename."""
    # Try magic first
    try:
        mime_type = magic.from_buffer(file_content, mime=True)
        if mime_type:
            return mime_type
    except:
        pass
    
    # Fallback to filename extension
    mime_type, _ = mimetypes.guess_type(filename)
    return mime_type or 'application/octet-stream'


def extract_text_from_pdf(file_content: bytes) -> Tuple[str, float]:
    """Extract text from PDF file."""
    try:
        text = extract_pdf_text(io.BytesIO(file_content))
        confidence = 1.0 if text and text.strip() else 0.0
        return text.strip(), confidence
    except Exception as e:
        print(f"PDF extraction failed: {e}")
        return "", 0.0


def extract_text_from_docx(file_content: bytes) -> Tuple[str, float]:
    """Extract text from DOCX file."""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        
        confidence = 1.0 if text and text.strip() else 0.0
        return text.strip(), confidence
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return "", 0.0


def extract_text(file_path: str) -> Tuple[str, float]:
    """Extract text from local file."""
    try:
        # Read file content
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        # Detect file type
        file_type = detect_file_type(file_content, file_path)
        
        # Extract based on type
        if 'pdf' in file_type:
            return extract_text_from_pdf(file_content)
        elif 'word' in file_type or 'officedocument' in file_type:
            return extract_text_from_docx(file_content)
        else:
            # Try as plain text
            try:
                text = file_content.decode('utf-8')
                return text, 1.0
            except:
                return "Unsupported file type. Please upload PDF or DOCX files.", 0.0
                
    except Exception as e:
        print(f"Text extraction failed for {file_path}: {e}")
        return "", 0.0
